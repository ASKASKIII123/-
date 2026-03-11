#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

try:
    from processors.word_processor import WordProcessor
    from processors.pdf_processor import PDFProcessor
except ImportError as e:
    print(f"错误: 无法导入处理器模块 - {e}")
    print("请确保已安装所有依赖包")
    sys.exit(1)

def print_header():
    print("=" * 60)
    print("   论文格式自动修改软件 - 命令行版本")
    print("=" * 60)
    print()

def print_menu():
    print("请选择操作:")
    print("1. 格式化Word文档")
    print("2. 格式化PDF文档")
    print("3. 创建测试文档")
    print("4. 退出")
    print()

def format_word_document():
    print("\n--- 格式化Word文档 ---")
    
    file_path = input("请输入Word文档路径 (.docx): ").strip()
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return
    
    print("\n选择模板:")
    print("1. 学术论文模板")
    print("2. 期刊论文模板")
    print("3. 学位论文模板")
    print("4. 会议论文模板")
    print("5. 自定义设置")
    
    template_choice = input("\n请选择 (1-5): ").strip()
    
    output_path = input("请输入输出文件路径 (留空则自动生成): ").strip()
    if not output_path:
        base_name = os.path.splitext(file_path)[0]
        output_path = f"{base_name}_formatted.docx"
    
    try:
        processor = WordProcessor(file_path)
        
        template_map = {
            "1": "academic",
            "2": "journal",
            "3": "thesis",
            "4": "conference"
        }
        
        if template_choice in template_map:
            template_name = template_map[template_choice]
            print(f"\n正在应用 {template_name} 模板...")
            processor.apply_template(template_name)
        elif template_choice == "5":
            print("\n自定义设置:")
            font_name = input("正文字体 (默认: 宋体): ").strip() or "宋体"
            font_size = int(input("正文字号 (默认: 12): ").strip() or "12")
            line_spacing = float(input("行间距 (默认: 1.5): ").strip() or "1.5")
            
            processor.set_font(font_name=font_name, font_size=font_size)
            processor.set_paragraph_format(line_spacing=line_spacing)
        
        processor.save(output_path)
        print(f"\n✓ 格式化完成！")
        print(f"输出文件: {output_path}")
        
    except Exception as e:
        print(f"\n✗ 格式化失败: {str(e)}")

def format_pdf_document():
    print("\n--- 格式化PDF文档 ---")
    
    file_path = input("请输入PDF文档路径 (.pdf): ").strip()
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return
    
    print("\n选择模板:")
    print("1. 学术论文模板")
    print("2. 期刊论文模板")
    print("3. 学位论文模板")
    print("4. 会议论文模板")
    
    template_choice = input("\n请选择 (1-4): ").strip()
    
    output_path = input("请输入输出文件路径 (留空则自动生成): ").strip()
    if not output_path:
        base_name = os.path.splitext(file_path)[0]
        output_path = f"{base_name}_formatted.pdf"
    
    try:
        processor = PDFProcessor(file_path)
        
        template_map = {
            "1": "academic",
            "2": "journal",
            "3": "thesis",
            "4": "conference"
        }
        
        if template_choice in template_map:
            template_name = template_map[template_choice]
            print(f"\n正在应用 {template_name} 模板...")
            processor.apply_template(template_name)
        
        processor.save(output_path)
        print(f"\n✓ 格式化完成！")
        print(f"输出文件: {output_path}")
        
    except Exception as e:
        print(f"\n✗ 格式化失败: {str(e)}")

def create_test_document():
    print("\n--- 创建测试文档 ---")
    
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        doc.add_heading('测试文档标题', 0)
        doc.add_heading('第一章 引言', level=1)
        doc.add_paragraph('这是一个测试文档，用于测试论文格式自动修改软件的功能。')
        doc.add_heading('1.1 背景', level=2)
        doc.add_paragraph('在学术写作中，格式的一致性非常重要。')
        doc.add_heading('第二章 方法', level=1)
        doc.add_paragraph('本章介绍测试的方法和步骤。')
        doc.add_heading('参考文献', level=1)
        doc.add_paragraph('[1] 张三. 论文格式研究[J]. 学术期刊, 2023, 10(2): 123-145.')
        doc.add_paragraph('[2] 李四. 文档排版技术[M]. 北京: 出版社, 2022.')
        
        output_path = "test_document.docx"
        doc.save(output_path)
        
        print(f"\n✓ 测试文档创建成功！")
        print(f"文件路径: {output_path}")
        print("\n现在可以使用此文档测试格式化功能")
        
    except Exception as e:
        print(f"\n✗ 创建测试文档失败: {str(e)}")
        print("请确保已安装 python-docx 包")

def main():
    print_header()
    
    while True:
        print_menu()
        choice = input("请选择 (1-4): ").strip()
        
        if choice == "1":
            format_word_document()
        elif choice == "2":
            format_pdf_document()
        elif choice == "3":
            create_test_document()
        elif choice == "4":
            print("\n感谢使用论文格式自动修改软件！")
            break
        else:
            print("\n无效的选择，请重新输入")
        
        input("\n按回车键继续...")
        print("\n" + "=" * 60 + "\n")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n程序被用户中断")
    except Exception as e:
        print(f"\n发生错误: {str(e)}")
        import traceback
        traceback.print_exc()