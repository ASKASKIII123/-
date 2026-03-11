#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt

def create_test_document():
    doc = Document()
    
    doc.add_heading('测试文档标题', 0)
    
    doc.add_heading('第一章 引言', level=1)
    doc.add_paragraph('这是一个测试文档，用于测试论文格式自动修改软件的功能。')
    doc.add_paragraph('本文档包含了各种格式的内容，包括标题、正文、引用等。')
    
    doc.add_heading('1.1 背景', level=2)
    doc.add_paragraph('在学术写作中，格式的一致性非常重要。')
    doc.add_paragraph('正确的格式可以提高文档的可读性和专业性。')
    
    doc.add_heading('1.2 目的', level=2)
    doc.add_paragraph('本文档的目的是测试格式化软件的各项功能。')
    doc.add_paragraph('包括字体设置、段落格式、页面布局等。')
    
    doc.add_heading('第二章 方法', level=1)
    doc.add_paragraph('本章介绍测试的方法和步骤。')
    
    doc.add_heading('2.1 测试步骤', level=2)
    p = doc.add_paragraph('测试步骤如下：')
    p.add_run('\n1. 选择测试文档')
    p.add_run('\n2. 应用预设模板')
    p.add_run('\n3. 自定义格式设置')
    p.add_run('\n4. 执行格式化')
    p.add_run('\n5. 检查格式化结果')
    
    doc.add_heading('第三章 结果', level=1)
    doc.add_paragraph('测试结果显示软件能够正确格式化文档。')
    
    doc.add_heading('第四章 结论', level=1)
    doc.add_paragraph('论文格式自动修改软件是一个实用的工具。')
    doc.add_paragraph('它可以帮助用户快速统一文档格式。')
    
    doc.add_heading('参考文献', level=1)
    doc.add_paragraph('[1] 张三. 论文格式研究[J]. 学术期刊, 2023, 10(2): 123-145.')
    doc.add_paragraph('[2] 李四. 文档排版技术[M]. 北京: 出版社, 2022.')
    doc.add_paragraph('[3] 王五. 自动化工具开发[D]. 某大学, 2021.')
    
    return doc

if __name__ == "__main__":
    doc = create_test_document()
    doc.save('test_document.docx')
    print("测试文档已创建: test_document.docx")