#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import copy

class WordProcessor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = Document(file_path)
    
    def save(self, output_path=None):
        if output_path is None:
            output_path = self.file_path
        self.doc.save(output_path)
    
    def set_font(self, font_name="宋体", font_size=12, bold=False, italic=False):
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run.font.italic = italic
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def set_title_font(self, font_name="黑体", font_size=16, bold=True):
        for paragraph in self.doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.bold = bold
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def set_paragraph_format(self, line_spacing=1.5, first_line_indent=2, space_before=0, space_after=0):
        for paragraph in self.doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = line_spacing
            paragraph_format.first_line_indent = Cm(first_line_indent)
            paragraph_format.space_before = Pt(space_before)
            paragraph_format.space_after = Pt(space_after)
    
    def set_page_margins(self, top=2.54, bottom=2.54, left=3.17, right=3.17):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(top)
            section.bottom_margin = Cm(bottom)
            section.left_margin = Cm(left)
            section.right_margin = Cm(right)
    
    def set_page_size(self, width=21.0, height=29.7):
        sections = self.doc.sections
        for section in sections:
            section.page_width = Cm(width)
            section.page_height = Cm(height)
    
    def set_alignment(self, alignment="center"):
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        for paragraph in self.doc.paragraphs:
            paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    def format_references_apa(self):
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith('[') and ']' in text:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.hanging_indent = Cm(1.27)
    
    def format_references_mla(self):
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith('[') and ']' in text:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.hanging_indent = Cm(1.27)
    
    def format_references_cssci(self):
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith('[') and ']' in text:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.hanging_indent = Cm(1.27)
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    
    def apply_template(self, template_name):
        templates = {
            "academic": self._apply_academic_template,
            "journal": self._apply_journal_template,
            "thesis": self._apply_thesis_template,
            "conference": self._apply_conference_template,
            "imrad": self._apply_imrad_template,
            "cssci": self._apply_cssci_template
        }
        
        if template_name in templates:
            templates[template_name]()
    
    def _apply_academic_template(self):
        self.set_font(font_name="宋体", font_size=12)
        self.set_title_font(font_name="黑体", font_size=16, bold=True)
        self.set_paragraph_format(line_spacing=1.5, first_line_indent=2)
        self.set_page_margins(top=2.54, bottom=2.54, left=3.17, right=3.17)
    
    def _apply_journal_template(self):
        self.set_font(font_name="Times New Roman", font_size=11)
        self.set_title_font(font_name="Arial", font_size=14, bold=True)
        self.set_paragraph_format(line_spacing=2.0, first_line_indent=1.27)
        self.set_page_margins(top=2.54, bottom=2.54, left=2.54, right=2.54)
    
    def _apply_thesis_template(self):
        self.set_font(font_name="宋体", font_size=14)
        self.set_title_font(font_name="黑体", font_size=18, bold=True)
        self.set_paragraph_format(line_spacing=1.5, first_line_indent=2)
        self.set_page_margins(top=3.0, bottom=3.0, left=3.0, right=3.0)
    
    def _apply_conference_template(self):
        self.set_font(font_name="宋体", font_size=10)
        self.set_title_font(font_name="黑体", font_size=12, bold=True)
        self.set_paragraph_format(line_spacing=1.25, first_line_indent=2)
        self.set_page_margins(top=2.0, bottom=2.0, left=2.5, right=2.5)
    
    def _apply_imrad_template(self):
        self.set_font(font_name="Times New Roman", font_size=12)
        self.set_title_font(font_name="Arial", font_size=14, bold=True)
        self.set_paragraph_format(line_spacing=1.5, first_line_indent=1.27)
        self.set_page_margins(top=2.54, bottom=2.54, left=2.54, right=2.54)
    
    def _apply_cssci_template(self):
        self.set_font(font_name="宋体", font_size=12)
        self.set_title_font(font_name="黑体", font_size=16, bold=True)
        self.set_paragraph_format(line_spacing=1.5, first_line_indent=2)
        self.set_page_margins(top=2.54, bottom=2.54, left=3.17, right=3.17)
        self.format_references_cssci()
    
    def enable_comparison_mode(self, original_file_path):
        original_doc = Document(original_file_path)
        
        for i, (orig_para, new_para) in enumerate(zip(original_doc.paragraphs, self.doc.paragraphs)):
            if orig_para.text != new_para.text:
                self._highlight_paragraph_changes(new_para, orig_para.text)
        
        for i, (orig_run, new_run) in enumerate(zip(original_doc.paragraphs, self.doc.paragraphs)):
            if orig_run.text != new_run.text:
                self._highlight_text_changes(new_run, orig_run.text)
    
    def _highlight_paragraph_changes(self, paragraph, original_text):
        if paragraph.text and not original_text:
            self._add_highlight(paragraph, "green")
        elif not paragraph.text and original_text:
            self._add_strikethrough(paragraph, "red")
    
    def _highlight_text_changes(self, paragraph, original_text):
        if not paragraph.runs:
            return
        
        new_text = paragraph.text
        orig_text = original_text
        
        if new_text and not orig_text:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 128, 0)
        elif not new_text and orig_text:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.strike = True
    
    def _add_highlight(self, paragraph, color):
        for run in paragraph.runs:
            if color == "green":
                run.font.color.rgb = RGBColor(0, 128, 0)
            elif color == "red":
                run.font.color.rgb = RGBColor(255, 0, 0)
    
    def _add_strikethrough(self, paragraph, color):
        for run in paragraph.runs:
            run.font.strike = True
            if color == "red":
                run.font.color.rgb = RGBColor(255, 0, 0)
    
    def get_chapters(self):
        chapters = []
        current_chapter = None
        chapter_keywords = ['摘要', '引言', '方法', '结果', '讨论', '结论', '参考文献', 
                        'Abstract', 'Introduction', 'Methods', 'Results', 'Discussion', 
                        'Conclusion', 'References', '第一章', '第二章', '第三章', 
                        '第四章', '第五章', '第六章', 'Chapter']
        
        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text.strip()
            if any(keyword in text for keyword in chapter_keywords):
                current_chapter = {
                    'name': text,
                    'start_index': i,
                    'paragraphs': [paragraph]
                }
                chapters.append(current_chapter)
            elif current_chapter:
                current_chapter['paragraphs'].append(paragraph)
        
        return chapters
    
    def format_chapter(self, chapter_index, template_name=None, font_name=None, font_size=None, 
                     line_spacing=None, first_line_indent=None):
        chapters = self.get_chapters()
        
        if chapter_index >= len(chapters):
            raise IndexError(f"章节索引 {chapter_index} 超出范围，共有 {len(chapters)} 个章节")
        
        chapter = chapters[chapter_index]
        
        for paragraph in chapter['paragraphs']:
            if template_name:
                self._apply_template_to_paragraph(paragraph, template_name)
            if font_name and font_size:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if line_spacing is not None:
                paragraph.paragraph_format.line_spacing = line_spacing
            if first_line_indent is not None:
                paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
    
    def _apply_template_to_paragraph(self, paragraph, template_name):
        templates = {
            "academic": {"font": "宋体", "size": 12, "line_spacing": 1.5, "indent": 2},
            "journal": {"font": "Times New Roman", "size": 11, "line_spacing": 2.0, "indent": 1.27},
            "thesis": {"font": "宋体", "size": 14, "line_spacing": 1.5, "indent": 2},
            "conference": {"font": "宋体", "size": 10, "line_spacing": 1.25, "indent": 2},
            "imrad": {"font": "Times New Roman", "size": 12, "line_spacing": 1.5, "indent": 1.27},
            "cssci": {"font": "宋体", "size": 12, "line_spacing": 1.5, "indent": 2}
        }
        
        if template_name in templates:
            tmpl = templates[template_name]
            for run in paragraph.runs:
                run.font.name = tmpl["font"]
                run.font.size = Pt(tmpl["size"])
                run._element.rPr.rFonts.set(qn('w:eastAsia'), tmpl["font"])
            paragraph.paragraph_format.line_spacing = tmpl["line_spacing"]
            paragraph.paragraph_format.first_line_indent = Cm(tmpl["indent"])
    
    def get_document_info(self):
        return {
            "paragraphs": len(self.doc.paragraphs),
            "pages": len(self.doc.sections),
            "file_size": os.path.getsize(self.file_path),
            "chapters": len(self.get_chapters())
        }